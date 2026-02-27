import streamlit as st
import fitz  # PyMuPDF
from optik1 import BubbleSheetScanner
import cv2
import numpy as np
import base64
import io
import os
import hashlib
import pickle
from mistralai import Mistral
from PIL import Image
from langchain_groq import ChatGroq
from sentence_transformers import SentenceTransformer
import faiss
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import json
import re

# ======================
# ⚙️  CONFIG
# ======================
REFERENCE_PDF_PATHS = [
    "references\الفيزياء 1-3.pdf",
    "references\الفيزياء 2.pdf",
    "references\الفيزياء 2-3.pdf"
]

FAISS_INDEX_DIR = "references/faiss_index"
os.makedirs(FAISS_INDEX_DIR, exist_ok=True)

TOP_K         = 6
CHUNK_SIZE    = 800
CHUNK_OVERLAP = 100

INDEX_FILE  = os.path.join(FAISS_INDEX_DIR, "index.faiss")
CHUNKS_FILE = os.path.join(FAISS_INDEX_DIR, "chunks.pkl")
HASH_FILE   = os.path.join(FAISS_INDEX_DIR, "hash.txt")
BUILT_FLAG  = os.path.join(FAISS_INDEX_DIR, "built.flag")

# ======================
# CLIENTS & MODELS
# ======================
st.set_page_config(page_title="Bubble Sheet Student Feedback", layout="wide")

API_KEY = st.secrets.get("GROQ_API_KEY")
llm = ChatGroq(
    model="llama-3.1-8b-instant",
    api_key=API_KEY,
    verbose=True,
    temperature=0.0,
)

MISTRAL_API_KEY = st.secrets.get("MISTRAL_API_KEY")
mistral_client = Mistral(api_key=MISTRAL_API_KEY)

scanner = BubbleSheetScanner(bubble_count=5)


# ======================
# HELPERS
# ======================

def compute_pdfs_hash(paths: list[str]) -> str:
    h = hashlib.md5()
    for p in paths:
        if os.path.exists(p):
            stat = os.stat(p)
            h.update(f"{p}:{stat.st_mtime}:{stat.st_size}".encode())
    return h.hexdigest()


def extract_text_from_pdf(path: str) -> str:
    doc = fitz.open(path)
    pages_text = [page.get_text("text") for page in doc]
    doc.close()
    return "\n".join(pages_text)


def chunk_text(text: str, source: str) -> list[dict]:
    chunks, start = [], 0
    while start < len(text):
        end   = min(start + CHUNK_SIZE, len(text))
        chunk = text[start:end].strip()
        if chunk:
            chunks.append({"text": chunk, "source": source})
        start += CHUNK_SIZE - CHUNK_OVERLAP
    return chunks


def clean_markdown(md: str) -> str:
    text = md.replace("- ", "").strip()
    text = text.replace("*", "").strip()
    return text


def index_is_ready() -> bool:
    return (
        os.path.exists(BUILT_FLAG)
        and os.path.exists(INDEX_FILE)
        and os.path.exists(CHUNKS_FILE)
    )


# ======================
# RAG — build once, load forever
# ======================

@st.cache_resource(show_spinner="📚 Loading / building RAG index...")
def get_rag_index(_pdf_paths: list[str]):
    embed_model = SentenceTransformer("all-MiniLM-L6-v2")

    if index_is_ready():
        index = faiss.read_index(INDEX_FILE)
        with open(CHUNKS_FILE, "rb") as f:
            chunks = pickle.load(f)
        pdfs_present = any(os.path.exists(p) for p in _pdf_paths)
        if pdfs_present and os.path.exists(HASH_FILE):
            current_hash = compute_pdfs_hash(_pdf_paths)
            with open(HASH_FILE) as f:
                saved_hash = f.read().strip()
            if saved_hash == current_hash:
                return index, chunks, embed_model
            else:
                st.info("🔄 PDFs changed — rebuilding index...")
        else:
            return index, chunks, embed_model

    available_pdfs = [p for p in _pdf_paths if os.path.exists(p)]
    if not available_pdfs:
        st.error(
            "❌ No pre-built index found and no reference PDFs available.\n"
            "Add the PDFs to `references/` once to build the index, then you can delete them."
        )
        return None, [], embed_model

    all_chunks = []
    for path in available_pdfs:
        text = extract_text_from_pdf(path)
        all_chunks.extend(chunk_text(text, source=os.path.basename(path)))

    if not all_chunks:
        st.error("❌ No text extracted from PDFs.")
        return None, [], embed_model

    texts = [c["text"] for c in all_chunks]
    BATCH = 256
    embeddings = np.vstack([
        embed_model.encode(texts[i:i + BATCH], show_progress_bar=False)
        for i in range(0, len(texts), BATCH)
    ]).astype("float32")

    index = faiss.IndexFlatL2(embeddings.shape[1])
    index.add(embeddings)

    faiss.write_index(index, INDEX_FILE)
    with open(CHUNKS_FILE, "wb") as f:
        pickle.dump(all_chunks, f)
    with open(HASH_FILE, "w") as f:
        f.write(compute_pdfs_hash(available_pdfs))
    with open(BUILT_FLAG, "w") as f:
        f.write("ok")

    st.success(
        f"✅ Index built from {len(available_pdfs)} PDF(s) — "
        f"{len(all_chunks):,} chunks saved. You may now delete the PDF files."
    )
    return index, all_chunks, embed_model


def retrieve(query: str, index, chunks: list[dict], embed_model) -> str:
    if index is None or not chunks:
        return ""
    query_vec = embed_model.encode([query], show_progress_bar=False).astype("float32")
    _, indices = index.search(query_vec, TOP_K)
    results = [
        f"[{chunks[i]['source']}]\n{chunks[i]['text']}"
        for i in indices[0] if 0 <= i < len(chunks)
    ]
    return "\n\n---\n\n".join(results)


# ======================
# Bubble Sheet Logic
# ======================

def process_bubble_sheet(image):
    h = int(600 * image.shape[0] / image.shape[1])
    frame = cv2.resize(image, (600, h))
    canny  = scanner.getCannyFrame(frame)
    warped = scanner.getWarpedFrame(canny, frame)
    if warped is None:
        raise ValueError("Bubble sheet not detected")

    adaptive = scanner.getAdaptiveThresh(frame)
    ovals    = scanner.getOvalContours(adaptive)
    q_count  = len(ovals) // scanner.bubbleCount
    ovals    = sorted(ovals, key=scanner.y_cord)
    answers, fill_threshold = [], 1.0

    for q in range(q_count):
        bubbles = sorted(
            ovals[q * scanner.bubbleCount:(q + 1) * scanner.bubbleCount],
            key=scanner.x_cord,
        )
        best_idx, best_ratio = None, 0
        for j, c in enumerate(bubbles):
            area   = cv2.contourArea(c)
            mask   = np.zeros(adaptive.shape, dtype="uint8")
            cv2.drawContours(mask, [c], -1, 255, -1)
            filled = cv2.countNonZero(cv2.bitwise_and(adaptive, adaptive, mask=mask))
            ratio  = filled / area if area else 0
            if ratio > best_ratio:
                best_ratio, best_idx = ratio, j
        answers.append(best_idx if best_ratio > fill_threshold else None)

    return answers, warped


# ======================
# LLM: Auto-solve exam questions → correct answers
# ======================

def llm_solve_exam(questions_text: str, rag_context: str, num_questions: int) -> list:
    """
    Ask the LLM to read the MCQ questions and return the correct answer
    for each question as a JSON list like ["A","B","C","D",...].
    Falls back to empty list on parse failure.
    """
    prompt = f"""
أنت معلم خبير في الفيزياء للمرحلة الثانوية.

فيما يلي مقتطفات من المراجع الدراسية:
─────────────────────────────
{rag_context or "لا توجد مراجع."}
─────────────────────────────

وفيما يلي أسئلة امتحان الاختيار من متعدد:
─────────────────────────────
{questions_text}
─────────────────────────────

المطلوب:
- اقرأ كل سؤال بعناية واختر الإجابة الصحيحة من الخيارات المتاحة.
- عدد الأسئلة المتوقع: {num_questions}
- أعد النتيجة فقط كـ JSON array بهذا الشكل بالضبط، بدون أي نص إضافي:
["A","B","C","D","A",...]

مهم جداً:
- لا تكتب أي شرح أو مقدمة أو خاتمة.
- فقط الـ JSON array.
- إذا لم تستطع تحديد إجابة لسؤال ما، ضع null.
"""
    raw = llm.invoke([{"role": "user", "content": prompt}]).content.strip()

    # Extract JSON array robustly
    match = re.search(r'\[.*?\]', raw, re.DOTALL)
    if not match:
        return []

    try:
        arr = json.loads(match.group())
    except json.JSONDecodeError:
        return []

    mapping = {"A": 0, "B": 1, "C": 2, "D": 3, "E": 4}
    return [mapping.get(str(x).upper()) if x is not None else None for x in arr]


# ======================
# Answer Comparison
# ======================

def compare_answers(student_answers: list, correct_answers: list) -> list[dict]:
    results = []
    max_q = max(len(student_answers), len(correct_answers))
    for i in range(max_q):
        s = student_answers[i] if i < len(student_answers) else None
        c = correct_answers[i] if i < len(correct_answers) else None
        s_letter = chr(ord("A") + s) if s is not None else "فارغ"
        c_letter = chr(ord("A") + c) if c is not None else "—"
        is_correct = (s == c) if (s is not None and c is not None) else False
        results.append({
            "question": i + 1,
            "student": s_letter,
            "correct": c_letter,
            "is_correct": is_correct,
        })
    return results


def score_summary(comparison: list[dict]) -> tuple[int, int]:
    total   = len(comparison)
    correct = sum(1 for r in comparison if r["is_correct"])
    return correct, total


# ======================
# LLM Full Analysis
# ======================

def analyze_student_performance(
    questions_text: str,
    comparison: list[dict],
    rag_context: str,
    student_name: str = "",
    student_class: str = "",
    semester: str = "",
) -> str:
    correct_count, total = score_summary(comparison)
    percentage = round((correct_count / total) * 100, 1) if total > 0 else 0

    comparison_text = "\n".join([
        f"سؤال {r['question']}: إجابة الطالب = {r['student']} | الإجابة الصحيحة = {r['correct']} | {'✅ صحيح' if r['is_correct'] else '❌ خطأ'}"
        for r in comparison
    ])

    student_info_lines = []
    if student_name:
        student_info_lines.append(f"• الاسم: {student_name}")
    if student_class:
        student_info_lines.append(f"• الصف: {student_class}")
    if semester:
        student_info_lines.append(f"• الفصل الدراسي: {semester}")
    student_info_block = "\n".join(student_info_lines) if student_info_lines else "لم يتم إدخال بيانات الطالب."

    prompt = f"""
أنت معلم خبير في تدريس الفيزياء والعلوم للمرحلة الثانوية.

─────────────────────────────
👤 بيانات الطالب:
─────────────────────────────
{student_info_block}

─────────────────────────────
📊 نتيجة المقارنة: {correct_count} صحيح من أصل {total} سؤال ({percentage}%)
─────────────────────────────
{comparison_text}

─────────────────────────────
📚 مقتطفات من المراجع الدراسية المعتمدة:
─────────────────────────────
{rag_context or "لا توجد مراجع متاحة."}

─────────────────────────────
📝 أسئلة الامتحان:
─────────────────────────────
{questions_text}
─────────────────────────────

بناءً على نتائج المقارنة والمراجع الدراسية، اكتب تقريراً شاملاً يتضمن الأقسام التالية:

1. **ملخص الأداء**: الدرجة والنسبة المئوية والتقييم العام.
2. **تحليل الأخطاء**: لكل سؤال خاطئ، اشرح المفهوم الصحيح باختصار.
3. **الفصول والوحدات المطلوب مراجعتها**: حدد بدقة أسماء الفصول والوحدات من المراجع الواردة في المقتطفات فقط — لا تذكر مصادر خارجية.
4. **خطة المراجعة**: خطوات عملية منظمة مبنية على الفصول المحددة.
5. **نقاط القوة**: المواضيع التي أتقنها الطالب.

اكتب بالعربية الفصحى بأسلوب تشجيعي وبنّاء.
لا تكرر الإجابات الحرفية — ركّز على المفاهيم والفهم.
"""
    return llm.invoke([{"role": "user", "content": prompt}]).content


# ======================
# Word Document Generation
# ======================

def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_heading_p(doc: Document, text: str, level: int = 1, color: RGBColor = None):
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16 - (level * 2))
    if color:
        run.font.color.rgb = color
    return p


def add_rtl_paragraph(doc: Document, text: str, bold: bool = False, size: int = 12):
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def add_comparison_table(doc: Document, comparison: list[dict]):
    add_heading_p(doc, "📊 جدول مقارنة الإجابات", level=2, color=RGBColor(0x1F, 0x49, 0x7D))

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    headers = ["رقم السؤال", "إجابة الطالب", "الإجابة الصحيحة", "النتيجة"]
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc_pr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F497D")
        shd.set(qn("w:color"), "FFFFFF")
        shd.set(qn("w:val"), "clear")
        tc_pr.append(shd)
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for r in comparison:
        row = table.add_row().cells
        row[0].text = str(r["question"])
        row[1].text = r["student"]
        row[2].text = r["correct"]
        row[3].text = "✅ صحيح" if r["is_correct"] else "❌ خطأ"
        for cell in row:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        fill_color = "E2EFDA" if r["is_correct"] else "FCE4D6"
        for cell in row:
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), fill_color)
            shd.set(qn("w:val"), "clear")
            tc_pr.append(shd)

    doc.add_paragraph()


def generate_word_report(
    student_name: str,
    student_class: str,
    semester: str,
    comparison: list[dict],
    feedback: str,
) -> bytes:
    doc = Document()

    section = doc.sections[0]
    section.page_width  = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = section.right_margin = Inches(1.0)

    doc.styles["Normal"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("تقرير أداء الطالب")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run("نظام تصحيح ورقة الإجابات والتغذية الراجعة الذكية")
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    doc.add_paragraph()

    # Student info
    add_heading_p(doc, "👤 بيانات الطالب", level=2, color=RGBColor(0x1F, 0x49, 0x7D))
    info_table = doc.add_table(rows=1, cols=3)
    info_table.style = "Table Grid"
    cells = info_table.rows[0].cells
    cells[0].text = f"الاسم: {student_name or '—'}"
    cells[1].text = f"الصف: {student_class or '—'}"
    cells[2].text = f"الفصل الدراسي: {semester or '—'}"
    for cell in cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].bold = True
    doc.add_paragraph()

    # Score
    correct_count, total = score_summary(comparison)
    percentage = round((correct_count / total) * 100, 1) if total > 0 else 0
    add_heading_p(doc, "🏆 ملخص النتيجة", level=2, color=RGBColor(0x1F, 0x49, 0x7D))
    score_p = doc.add_paragraph()
    set_rtl(score_p)
    score_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    score_run = score_p.add_run(f"{correct_count} / {total}   ({percentage}%)")
    score_run.bold = True
    score_run.font.size = Pt(18)
    score_run.font.color.rgb = (
        RGBColor(0x37, 0x86, 0x10) if percentage >= 80
        else RGBColor(0xFF, 0x8C, 0x00) if percentage >= 60
        else RGBColor(0xC0, 0x00, 0x00)
    )
    doc.add_paragraph()

    # Comparison table
    add_comparison_table(doc, comparison)

    # LLM Feedback
    add_heading_p(doc, "📋 التقرير التفصيلي والتوصيات", level=2, color=RGBColor(0x1F, 0x49, 0x7D))
    for line in feedback.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        is_heading_line = line.startswith(("**", "##", "#", "١.", "٢.", "٣.", "٤.", "٥.",
                                           "1.", "2.", "3.", "4.", "5."))
        clean_line = line.replace("**", "").replace("##", "").replace("#", "").strip()
        add_rtl_paragraph(doc, clean_line, bold=is_heading_line, size=12 if not is_heading_line else 13)

    # Footer
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run(
        f"تم إنشاء هذا التقرير بتاريخ: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ======================
# Streamlit UI
# ======================
st.title("📝 Bubble Sheet + Student Feedback (Arabic)")

faiss_index, all_chunks, embed_model = get_rag_index(REFERENCE_PDF_PATHS)

if faiss_index is not None:
    st.success(f"✅ RAG جاهز — {len(all_chunks):,} chunk من {len(REFERENCE_PDF_PATHS)} مراجع")
else:
    st.error("❌ فشل تحميل المراجع. تحقق من المسارات في REFERENCE_PDF_PATHS.")

st.divider()

# Student Information
st.subheader("👤 بيانات الطالب")
col1, col2, col3 = st.columns(3)
with col1:
    student_name = st.text_input("اسم الطالب", placeholder="أدخل اسم الطالب الكامل")
with col2:
    student_class = st.selectbox(
        "الصف",
        ["الصف الأول الثانوي", "الصف الثاني الثانوي", "الصف الثالث الثانوي"],
    )
with col3:
    semester = st.selectbox(
        "الفصل الدراسي",
        ["الفصل الأول", "الفصل الثاني"],
    )

st.divider()

# PDF Upload
uploaded_file = st.file_uploader("Upload Exam PDF", type=["pdf"])

if uploaded_file:
    with st.spinner("⚙️ جاري معالجة ورقة الإجابات..."):
        uploaded_bytes = uploaded_file.read()
        doc_pdf = fitz.open(stream=uploaded_bytes, filetype="pdf")
        answers, questions_text = [], ""

        for i in range(len(doc_pdf)):
            pix     = doc_pdf[i].get_pixmap(matrix=fitz.Matrix(300 / 72, 300 / 72))
            pil_img = Image.open(io.BytesIO(pix.tobytes("png")))
            if i == 0:
                cv_img = cv2.cvtColor(np.array(pil_img), cv2.COLOR_RGB2BGR)
                try:
                    answers, _ = process_bubble_sheet(cv_img)
                except ValueError:
                    st.warning("⚠️ لم يُكتشف ورقة الإجابة — سيُعامَل الامتحان كمقالي.")

        # OCR via Mistral
        b64_pdf  = base64.b64encode(uploaded_bytes).decode("utf-8")
        ocr_resp = mistral_client.ocr.process(
            model="mistral-ocr-latest",
            document={"type": "document_url", "document_url": f"data:application/pdf;base64,{b64_pdf}"},
            include_image_base64=True,
        )
        start_page = 1 if answers else 0
        for page in ocr_resp.pages[start_page:]:
            questions_text += page.markdown + "\n\n"

    # RAG retrieval
    rag_context = ""
    if faiss_index is not None and questions_text.strip():
        with st.spinner("🔍 جاري استرجاع المواد الدراسية ذات الصلة..."):
            rag_context = retrieve(questions_text[:1500], faiss_index, all_chunks, embed_model)
            rag_context = clean_markdown(rag_context)

    # ── LLM solves the exam automatically ─────────────────────
    correct_answers = []
    if answers and questions_text.strip():
        with st.spinner("🧠 الذكاء الاصطناعي يحل الأسئلة ويحدد الإجابات الصحيحة..."):
            correct_answers = llm_solve_exam(questions_text, rag_context, len(answers))

    # ── Answer Comparison ──────────────────────────────────────
    comparison = []
    if answers and correct_answers:
        comparison = compare_answers(answers, correct_answers)
        correct_count, total = score_summary(comparison)
        percentage = round((correct_count / total) * 100, 1) if total > 0 else 0

        st.subheader("📊 مقارنة إجابات الطالب بالإجابات الصحيحة")
        st.caption("✅ الإجابات الصحيحة تم تحديدها تلقائياً بواسطة الذكاء الاصطناعي")

        col_a, col_b, col_c = st.columns(3)
        col_a.metric("الدرجة", f"{correct_count} / {total}")
        col_b.metric("النسبة المئوية", f"{percentage}%")
        grade = (
            "ممتاز 🌟" if percentage >= 90
            else "جيد جداً ✅" if percentage >= 75
            else "جيد 👍" if percentage >= 60
            else "يحتاج مراجعة ⚠️"
        )
        col_c.metric("التقدير", grade)

        import pandas as pd
        df = pd.DataFrame([{
            "رقم السؤال": r["question"],
            "إجابة الطالب": r["student"],
            "الإجابة الصحيحة (AI)": r["correct"],
            "النتيجة": "✅" if r["is_correct"] else "❌",
        } for r in comparison])

        def highlight_row(row):
            color = "background-color: #e2efda" if row["النتيجة"] == "✅" else "background-color: #fce4d6"
            return [color] * len(row)

        st.dataframe(
            df.style.apply(highlight_row, axis=1),
            use_container_width=True,
            hide_index=True,
        )

    elif answers and not correct_answers:
        st.warning("⚠️ تعذّر على الذكاء الاصطناعي تحديد الإجابات الصحيحة. تأكد من وضوح نص الأسئلة في الـ PDF.")

    # ── Full LLM Feedback ──────────────────────────────────────
    if comparison:
        st.subheader("🧠 التقرير التفصيلي والتوصيات")
        with st.spinner("📝 جاري إنشاء التقرير التفصيلي..."):
            feedback = analyze_student_performance(
                questions_text,
                comparison,
                rag_context,
                student_name=student_name,
                student_class=student_class,
                semester=semester,
            )
            feedback = clean_markdown(feedback)

        st.text_area("📌 تقرير الطالب", feedback, height=500)

        with st.expander("🔎 المقتطفات المسترجعة من المراجع"):
            st.text(rag_context or "لا توجد مقتطفات.")

        # ── Word Download ──────────────────────────────────────
        st.divider()
        st.subheader("💾 تحميل التقرير كملف Word")

        with st.spinner("📄 جاري إنشاء ملف Word..."):
            word_bytes = generate_word_report(
                student_name=student_name,
                student_class=student_class,
                semester=semester,
                comparison=comparison,
                feedback=feedback,
            )

        filename = f"تقرير_{student_name or 'الطالب'}_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.docx"

        st.download_button(
            label="⬇️ تحميل تقرير Word",
            data=word_bytes,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            type="primary",
        )